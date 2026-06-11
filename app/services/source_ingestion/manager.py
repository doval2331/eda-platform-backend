from __future__ import annotations

import json
import mimetypes
import urllib.error
import urllib.parse
import urllib.request
import uuid
from dataclasses import dataclass, field
from io import BytesIO
from typing import Literal

import pandas as pd

from app.config import get_settings

NormalizedKind = Literal["tabular", "text"]

TABULAR_EXTENSIONS = {".csv", ".tsv", ".xlsx", ".xlsm", ".json", ".parquet"}
TEXT_EXTENSIONS = {".txt", ".md", ".docx", ".pdf"}
AUDIO_EXTENSIONS = {".mp3", ".mp4", ".mpeg", ".mpga", ".m4a", ".wav", ".webm", ".ogg", ".flac"}
SUPPORTED_EXTENSIONS = TABULAR_EXTENSIONS | TEXT_EXTENSIONS | AUDIO_EXTENSIONS


@dataclass(frozen=True)
class IngestedSource:
    normalized_kind: NormalizedKind
    original_format: str
    extraction_method: str
    dataframe: pd.DataFrame | None = None
    text: str | None = None
    metadata: dict = field(default_factory=dict)


def ingest_source(filename: str, content: bytes) -> IngestedSource:
    ext = _extension(filename)
    normalized_kind = detect_source_kind(filename)
    if ext in TABULAR_EXTENSIONS:
        df, method, metadata = _read_tabular(ext, content)
        return IngestedSource(
            normalized_kind=normalized_kind,
            original_format=ext.lstrip("."),
            extraction_method=method,
            dataframe=_clean_dataframe(df),
            metadata=metadata,
        )
    if ext in TEXT_EXTENSIONS:
        text, method, metadata = _read_text(ext, content)
        return IngestedSource(
            normalized_kind=normalized_kind,
            original_format=ext.lstrip("."),
            extraction_method=method,
            text=_clean_text(text),
            metadata=metadata,
        )

    text, method, metadata = _transcribe_audio(filename, ext, content)
    return IngestedSource(
        normalized_kind=normalized_kind,
        original_format=ext.lstrip("."),
        extraction_method=method,
        text=_clean_text(text),
        metadata=metadata,
    )


def detect_source_kind(filename: str) -> NormalizedKind:
    ext = _extension(filename)
    if ext in TABULAR_EXTENSIONS:
        return "tabular"
    if ext in TEXT_EXTENSIONS or ext in AUDIO_EXTENSIONS:
        return "text"
    raise ValueError(
        "Formato no soportado. Se admiten CSV, TSV, XLSX, XLSM, JSON, Parquet, "
        "TXT, MD, DOCX, PDF y audio MP3/WAV/M4A/WEBM/OGG/FLAC."
    )


def _extension(filename: str) -> str:
    name = (filename or "").strip().lower()
    if "." not in name:
        return ""
    return name[name.rfind(".") :]


def _read_tabular(ext: str, content: bytes) -> tuple[pd.DataFrame, str, dict]:
    stream = BytesIO(content)
    if ext == ".csv":
        return _read_csv_like(stream, sep=None), "pandas.read_csv", {}
    if ext == ".tsv":
        return _read_csv_like(stream, sep="\t"), "pandas.read_csv(tsv)", {}
    if ext in {".xlsx", ".xlsm"}:
        return _read_excel_workbook(content)
    if ext == ".json":
        return _read_json_table(content), "pandas.read_json/json_normalize", {}
    if ext == ".parquet":
        try:
            return pd.read_parquet(stream), "pandas.read_parquet", {}
        except ImportError as exc:
            raise ValueError(
                "Para leer Parquet instala pyarrow o fastparquet, o exporta el archivo a CSV."
            ) from exc
    raise ValueError("Formato tabular no soportado.")


def _read_csv_like(stream: BytesIO, *, sep: str | None) -> pd.DataFrame:
    last_error: Exception | None = None
    for encoding in ("utf-8-sig", "utf-8", "latin-1"):
        stream.seek(0)
        try:
            kwargs = {"encoding": encoding}
            if sep is None:
                kwargs["sep"] = None
                kwargs["engine"] = "python"
            else:
                kwargs["sep"] = sep
            return pd.read_csv(stream, **kwargs)
        except Exception as exc:  # pragma: no cover - last error is reported
            last_error = exc
    raise ValueError(f"No se pudo leer el archivo tabular: {last_error}")


def _read_excel_workbook(content: bytes) -> tuple[pd.DataFrame, str, dict]:
    try:
        sheets = pd.read_excel(BytesIO(content), sheet_name=None, engine="openpyxl")
    except ImportError as exc:
        raise ValueError("Para leer Excel instala openpyxl.") from exc
    except Exception as exc:
        raise ValueError(f"No se pudo leer el Excel: {exc}") from exc

    usable: list[tuple[str, pd.DataFrame, int]] = []
    for sheet_name, df in sheets.items():
        cleaned = _clean_dataframe(df)
        score = int(cleaned.shape[0] * max(1, cleaned.shape[1]))
        if not cleaned.empty:
            usable.append((str(sheet_name), cleaned, score))
    if not usable:
        raise ValueError("El Excel no contiene hojas con datos tabulares.")
    sheet_name, df, _score = max(usable, key=lambda item: item[2])
    return df, "pandas.read_excel(openpyxl)", {"sheet_name": sheet_name}


def _read_json_table(content: bytes) -> pd.DataFrame:
    try:
        payload = json.loads(content.decode("utf-8-sig"))
    except UnicodeDecodeError as exc:
        raise ValueError("El JSON debe estar codificado en UTF-8.") from exc
    except json.JSONDecodeError as exc:
        raise ValueError(f"No se pudo leer el JSON: {exc}") from exc

    if isinstance(payload, list):
        return pd.json_normalize(payload)
    if isinstance(payload, dict):
        list_items = [(k, v) for k, v in payload.items() if isinstance(v, list)]
        if list_items:
            _key, values = max(list_items, key=lambda item: len(item[1]))
            return pd.json_normalize(values)
        return pd.json_normalize(payload)
    raise ValueError("El JSON debe contener un objeto o una lista de objetos.")


def _read_text(ext: str, content: bytes) -> tuple[str, str, dict]:
    if ext in {".txt", ".md"}:
        return _decode_text(content), "text.decode", {}
    if ext == ".docx":
        return _read_docx(content)
    if ext == ".pdf":
        return _read_pdf(content)
    raise ValueError("Formato documental no soportado.")


def _read_docx(content: bytes) -> tuple[str, str, dict]:
    try:
        from docx import Document
    except ImportError as exc:
        raise ValueError("Para leer Word instala python-docx.") from exc

    try:
        document = Document(BytesIO(content))
    except Exception as exc:
        raise ValueError(f"No se pudo leer el Word: {exc}") from exc

    parts = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if values:
                parts.append(" | ".join(values))
    return "\n".join(parts), "python-docx", {"paragraphs": len(document.paragraphs)}


def _read_pdf(content: bytes) -> tuple[str, str, dict]:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise ValueError("Para leer PDF instala pypdf.") from exc

    try:
        reader = PdfReader(BytesIO(content))
        pages = [(page.extract_text() or "").strip() for page in reader.pages]
    except Exception as exc:
        raise ValueError(f"No se pudo leer el PDF: {exc}") from exc

    text = "\n\n".join(page for page in pages if page)
    return text, "pypdf", {"page_count": len(reader.pages)}


def _decode_text(content: bytes) -> str:
    last_error: UnicodeDecodeError | None = None
    for encoding in ("utf-8-sig", "utf-8", "latin-1"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError as exc:
            last_error = exc
    raise ValueError("No se pudo decodificar el texto.") from last_error


def _transcribe_audio(filename: str, ext: str, content: bytes) -> tuple[str, str, dict]:
    settings = get_settings()
    if not settings.llm_enabled:
        raise ValueError("Para transcribir audio activa LLM_ENABLED=true en el backend.")
    if not settings.llm_api_key and "api.openai.com" in settings.llm_api_base:
        raise ValueError("Para transcribir audio configura LLM_API_KEY.")

    boundary = f"----tfm-source-{uuid.uuid4().hex}"
    model = settings.llm_transcription_model
    body = _multipart_body(
        boundary=boundary,
        fields={"model": model},
        files={
            "file": (
                filename,
                mimetypes.types_map.get(ext, "application/octet-stream"),
                content,
            )
        },
    )
    url = settings.llm_api_base.rstrip("/") + "/audio/transcriptions"
    api_version = settings.llm_api_version.strip()
    if api_version:
        url = f"{url}?{urllib.parse.urlencode({'api-version': api_version})}"

    headers = {"Content-Type": f"multipart/form-data; boundary={boundary}"}
    if settings.llm_api_key:
        if "openai.azure.com" in settings.llm_api_base:
            headers["api-key"] = settings.llm_api_key
        else:
            headers["Authorization"] = f"Bearer {settings.llm_api_key}"

    request = urllib.request.Request(url, data=body, headers=headers, method="POST")
    try:
        with urllib.request.urlopen(request, timeout=settings.llm_timeout_seconds) as response:
            data = json.loads(response.read().decode("utf-8"))
    except urllib.error.HTTPError as exc:
        raise ValueError(f"El proveedor speech-to-text respondio HTTP {exc.code}.") from exc
    except (urllib.error.URLError, TimeoutError, ValueError, OSError) as exc:
        raise ValueError(f"No se pudo transcribir el audio: {exc}") from exc

    text = str(data.get("text") or "").strip()
    if not text:
        raise ValueError("La transcripcion de audio no devolvio texto.")
    return text, "openai.audio.transcriptions", {"transcription_model": model}


def _multipart_body(
    *,
    boundary: str,
    fields: dict[str, str],
    files: dict[str, tuple[str, str, bytes]],
) -> bytes:
    chunks: list[bytes] = []
    for name, value in fields.items():
        chunks.extend(
            [
                f"--{boundary}\r\n".encode(),
                f'Content-Disposition: form-data; name="{name}"\r\n\r\n'.encode(),
                str(value).encode("utf-8"),
                b"\r\n",
            ]
        )
    for name, (filename, content_type, content) in files.items():
        chunks.extend(
            [
                f"--{boundary}\r\n".encode(),
                (
                    f'Content-Disposition: form-data; name="{name}"; '
                    f'filename="{filename}"\r\n'
                ).encode(),
                f"Content-Type: {content_type}\r\n\r\n".encode(),
                content,
                b"\r\n",
            ]
        )
    chunks.append(f"--{boundary}--\r\n".encode())
    return b"".join(chunks)


def _clean_dataframe(df: pd.DataFrame) -> pd.DataFrame:
    cleaned = df.copy()
    cleaned = cleaned.dropna(axis=0, how="all").dropna(axis=1, how="all")
    cleaned.columns = [str(col).strip() or f"col_{idx + 1}" for idx, col in enumerate(cleaned.columns)]
    return cleaned.reset_index(drop=True)


def _clean_text(text: str) -> str:
    cleaned = "\n".join(line.strip() for line in text.splitlines())
    cleaned = "\n".join(line for line in cleaned.splitlines() if line)
    if not cleaned:
        raise ValueError("No se pudo extraer texto util del archivo.")
    return cleaned
